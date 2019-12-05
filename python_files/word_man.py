import docx
import shade as shade
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches

from docx.oxml.shared import OxmlElement, qn

document = Document()

document.add_heading('Debarment Check', 0)

p = document.add_paragraph('Prior to being invited to participate in development/authoring of a publication sponsored '
                           'by Genzyme/Sanofi, a debarment check must be completed for each US author.')

# run = p.add_run()
# tag = run._r
# fld = docx.oxml.shared.OxmlElement('w:fldChar')
# fld.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
# fldData = docx.oxml.shared.OxmlElement('w:fldData')
#
# fldData.text = '/////2UAAAAUAAYAQwBoAGUAYwBrADEAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAAA'
# fldData.set(docx.oxml.ns.qn('xml:space'), 'preserve')
# fld.append(fldData)
# tag.append(fld)
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True

# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')
#
# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )


records = (
    ('Author Name', ''),
    ('Name of Institution', ''),
    ('City, State', '')
)

debarment_list = (
    ('Office of Inspectors General LIst of Excluded Individuals.', ''),
    ('System for Award Management', ''),
    ('Office of Research Integrity', '')
)

table = document.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Information'
hdr_cells[1].text = 'Id'

for qty, id in records:
    row_cells = table.add_row().cells
    row_cells[0].text = qty
    row_cells[1].text = id

document.add_paragraph('\n')

b_table = document.add_table(rows=1, cols=2)
b_table.style = 'Table Grid'
b_hdr_cells = b_table.rows[0].cells
b_hdr_cells[0].text = 'Debarment List'
b_hdr_cells[1].text = 'Findings'

for d_list, findings in debarment_list:
    b_row_cells = b_table.add_row().cells
    b_row_cells[0].text = d_list
    b_row_cells[1].text = findings

document.add_paragraph('\nIf the potential author is listed on any of the above, they may not be invited to author a '
                       'publication sponsored by Genzyme or Sanofi; advise publication lead of findings of this '
                       'search.\n')
document.add_paragraph('Once the debarment check has been completed, upload this document to the appropriate record '
                       'in Datavision.\n')
completion = document.add_paragraph('Debarment check completed by:\n')
completion.add_run('Debarment').bold = True
document.add_paragraph('Date check completed:\n')

document.add_picture('screenshot.png', width=Inches(6))

# Set a cell background (shading) color to RGB D9D9D9.
a_cell_1 = table.cell(0, 0)
a_co = a_cell_1._tc.get_or_add_tcPr()
a_cell_2 = table.cell(0, 1)
a_ct = a_cell_2._tc.get_or_add_tcPr()

a_cell_color_1 = OxmlElement('w:shd')
a_cell_color_1.set(qn('w:fill'), '#94C167')

a_cell_color_2 = OxmlElement('w:shd')
a_cell_color_2.set(qn('w:fill'), '#94C167')

a_co.append(a_cell_color_1)
a_ct.append(a_cell_color_2)





b_cell_1 = b_table.cell(0, 0)
b_co = b_cell_1._tc.get_or_add_tcPr()
b_cell_2 = b_table.cell(0, 1)
b_ct = b_cell_2._tc.get_or_add_tcPr()

b_cell_color_1 = OxmlElement('w:shd')
b_cell_color_1.set(qn('w:fill'), '#94C167')

b_cell_color_2 = OxmlElement('w:shd')
b_cell_color_2.set(qn('w:fill'), '#94C167')

b_co.append(b_cell_color_1)
b_ct.append(b_cell_color_2)


document.add_page_break()

document.save('demo.docx')
