#!/bin/python
# ===========================================================#
## bot version 2. This is the silent or headless version. Runs all procceses in the background

import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import openpyxl
from openpyxl import load_workbook
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from urllib.request import urlopen as uReq
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup as soup
from docx.oxml.shared import OxmlElement, qn
import datetime
import random
from datetime import timedelta
import shutil
import time
import re


Root_path = 'C:\\Users\\pelbanon\\Desktop\\DebarmentBotPackage\\'

# change directory
os.chdir("C:\\Users\\pelbanon\\Desktop\\DebarmentBotPackage")
logf = open(Root_path + 'log_file.txt', 'w')

driver = ""

# define the name of the directory to be created
# directory structure for screenshots, files, and compeleted files
def directory_structure(iter_rand):
    output = 'C:\\Users\\pelbanon\\Desktop\\DebarmentBotPackage\\Outputs_' + iter_rand
    debarment_file_path = 'C:\\Users\\pelbanon\\Desktop\\DebarmentBotPackage\\Outputs_' + iter_rand + '\\Debarment_files_' + iter_rand
    screenshots_path = 'C:\\Users\\pelbanon\\Desktop\\DebarmentBotPackage\\Outputs_' + iter_rand + '\\Screenshots_' + iter_rand
    completed_file_path = 'C:\\Users\\pelbanon\\Desktop\\DebarmentBotPackage\\Outputs_' + iter_rand + '\\Completed_files_' + iter_rand
    flagged_authors = 'C:\\Users\\pelbanon\\Desktop\\DebarmentBotPackage\\Outputs_' + iter_rand + '\\Flagged_Authors_' + iter_rand

    # error handlers for creating directory for screenshots and debarment file

    # if os.path.exists(output) is False:
    try:
        os.mkdir(output)

    except OSError as e:
        print("Creation of the directory %s failed" % output)
        logf.write('creation of the directory {0}:{1} failed\n'.format(str(output), str(e)))
        # print("Creation of the directory %s failed" % screenshots_path)
    else:
        print("Successfully created the directory %s " % output)

    try:
        os.mkdir(flagged_authors)
    except OSError as e:
        print("Creation of the directory %s failed" % flagged_authors)
        logf.write('creation of the directory {0}:{1} failed\n'.format(str(flagged_authors), str(e)))
        # print("Creation of the directory %s failed" % screenshots_path)
    else:
        print("Successfully created the directory %s " % flagged_authors)

    try:
        os.mkdir(screenshots_path)
    except OSError as e:
        print(e)
        logf.write('creation of the directory {0}:{1} failed\n'.format(str(screenshots_path), str(e)))
        # print("Creation of the directory %s failed" % screenshots_path)
    else:
        print("Successfully created the directory %s " % screenshots_path)

    try:
        os.mkdir(debarment_file_path)
    except OSError as e:

        print("Creation of the directory %s failed" % debarment_file_path)
        logf.write('creation of the directory {0}:{1} failed\n'.format(str(debarment_file_path), str(e)))
    else:
        print("Successfully created the directory %s " % debarment_file_path)

    try:
        os.mkdir(completed_file_path)
    except OSError as e:
        print("Creation of the directory %s failed" % completed_file_path)
        logf.write('creation of the directory {0}:{1} failed\n'.format(str(completed_file_path), str(e)))
    else:
        print("Successfully created the directory %s " % completed_file_path)


# gets the input variable for the process execution (excel file, list of authors)
def get_working_filename():
    if len(os.listdir('.\\working_dir')) == 0:
        logf.write('The directory is empty, please upload a file to the working directory\n')
        print('Error: The directory is empty, please upload a file to the working directory\n')
    elif len(os.listdir('.\\working_dir')) > 1:
        logf.write('Too many files, please upload one execel file at a time\n')
        print('Error: Too many files, please upload one execel file at a time\n')
    else:
        for file in os.listdir('.\\working_dir'):  # iterate through all the files
            if file.endswith('.xlsx'):  # identify if the file ends with.xlsx
                move_file = os.path.join(".\\working_dir", file)  # gets the file in working directory
                return move_file
            else:  # if file not ends with .xlsx then its not an excel file
                logf.write('Not an excel file\n')
                print('Error: The file is not an excel file. Please upload an excel file\n')


def check_not_missed(iter_rand):
    array = []
    for file in os.listdir(
            '.\\Outputs_' + iter_rand + '\\Debarment_files_' + iter_rand):  # iterate through all the files
        split_filename = file.split('_')
        array.append(split_filename[0])  # gets only the last name

    for flagged_file in os.listdir(
            '.\\Outputs_' + iter_rand + '\\Flagged_Authors_' + iter_rand):  # iterate through all the files
        split_filename = flagged_file.split('_')
        array.append(split_filename[0])  # gets only the last name

    return array


# get file and move file
def mv_dir_structure(iter_rand):
    try:
        # move the current file to the completed directory
        shutil.move(ex_name, ".\\Outputs_" + iter_rand + "\\Completed_files_" + iter_rand)
        print(f'Task: Moving {ex} to the completed folder')
    except shutil.Error as err:
        logf.write(str(err) + '\n')
        print(str(err))



try:
    # get webdriver for chrome chromedriver.exe path - this would have to change for everyone
    # driver = webdriver.Chrome(
    # 'Y:\\LINKS\\#LINKS Initiatives\\AI\\Debarment Checks\\DebarmentCheckResults\\chromedriver_win32\\chromedriver.exe')

    chrome_options = Options()
    # chrome_options.add_argument("--start-maximized")
    chrome_options.add_argument(f'window-size={1024}x{600}')
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--hide-scrollbars")


    driver = webdriver.Chrome(chrome_options=chrome_options,
                              executable_path='Y:\\LINKS\\#LINKS Initiatives\\AI\\Debarment Checks\\DebarmentCheckResults\\chromedriver_win32\\chromedriver.exe')

    # driver = webdriver.Chrome(chrome_options=chrome_options,
    #                           executable_path='C:\\Users\\yangb\\Downloads\\chromedriver_win32_test\\chromedriver.exe')

    print('Task: Chrome successfully opened')

except Exception as err:
    # driver.close()
    # driver.quit()
    print("error in driver")
    logf.write(str(err))

# sets the window size. This size is specified from screenshot reasons
# driver.set_window_size(1100, 1500)  # ideal was 1100, 1500
done = True  # flagger variable for flagging the process
# getting the file from the working_dir
ex_name = get_working_filename()  # excel file name (database)

ex = load_workbook(ex_name)  # opens the excel
# sheet =
sheet = ex[ex.sheetnames[0]]  # gets the first tab sheet

# url 1 for search
r_url = 'https://exclusions.oig.hhs.gov/'


# new sheet for not listed people
def create_sheet():
    if 'Flagged' not in ex.sheetnames:
        ex.create_sheet('Flagged')
        ex.save(ex_name)
    b_sheet = ex['Flagged']
    return b_sheet


def create_sheet_c():
    if 'Missed' not in ex.sheetnames:
        ex.create_sheet('Missed')
        ex.save(ex_name)
    c_sheet = ex['Missed']
    return c_sheet


# get name of one row
# same function
def get_name(r):
    # initializing bad_chars_list
    bad_chars = [';', '[', '@', '_', '!', '#', '$', '%', '^', '&', '*', '(', ')', '<', '>', '?', '/', '\\',
                 '|', '}', '{', '~', ':', ']']
    regex = re.compile("[@_!#$%^&*()<>?/\|}{~:]")

    # get the first name
    firstname = str(r[1].value)
    # split the first name into an array
    firstname_array = firstname.split()
    # if an array has more than 1 string then do the checking
    if len(firstname_array) > 1:
        counter = 1
        # if the length of the second string is 1 then authomatically remove
        if len(firstname_array[counter]) == 1:
            firstname = firstname_array[0]  # removed the middle name

            lastname = str(r[0].value).replace(" ", "")

            return lastname, firstname
        # else if the length is 2 then identify if there is a period
        elif len(firstname_array[counter]) == 2:
            string_name = firstname_array[counter]
            # if there is a period then remove the whole thing
            if "." in string_name:
                firstname = firstname_array[0]

                lastname = str(r[0].value).replace(" ", "")

                return lastname, firstname
            else:
                if regex.search(firstname_array[counter]) is None:
                    firstname = firstname_array[0] + " " + firstname_array[1]
                else:

                    for i in bad_chars:
                        firstname = firstname.replace(i, '')

                lastname = str(r[0].value).replace(" ", "")

                return lastname, firstname

        elif len(firstname_array[counter]) > 2:
            # firstname = firstname_array[0] + " " + firstname_array[1]
            # print(firstname)
            # lastname = str(r[0].value).replace(" ", "")
            # print(lastname)
            if regex.search(firstname_array[counter]) is None:
                firstname = firstname_array[0] + " " + firstname_array[1]


            else:
                firstname = firstname_array[0].replace(" ", "")
                # for i in bad_chars:
                #     firstname = firstname.replace(i, '')

            lastname = str(r[0].value).replace(" ", "")

            return lastname, firstname

    else:

        if regex.search(firstname) is None:
            firstname = firstname


        else:
            for i in bad_chars:
                firstname = firstname.replace(i, '')

        lastname = str(r[0].value).replace(" ", "")

        # firstname = str(r[1].value).replace(" ", "")
        # print(firstname)
        # lastname = str(r[0].value).replace(" ", "")
        # print(lastname)

        return lastname, firstname

        # else if there is no period but only one letther then remove the whole thing
        # else if no period then consider it as part of the first name
        # else if the len is more than 2 then consider it as part of the first name
    # else remove all spaces


# gets institution
def get_institution(r, row_number):
    if r[4].value is None:
        logf.write('institution is empty in row ' + row_number + '\n')
    else:
        institute = str(r[4].value)
    return institute


# gets the city state
def get_city(r, row_number):
    if r[5].value is None:
        logf.write('institution is empty in row ' + row_number + '\n')
        location = ' '
    else:
        location = str(r[5].value)
    return location

# gets the city state
def get_state(r, row_number):
    if r[6].value is None:
        logf.write('institution is empty in row ' + row_number + '\n')
        state = ' '
    else:
        state = str(r[6].value)
    return state


# gets the name of the contributer
def get_contributer(r, row_number):
    if r[8].value is None:
        logf.write('contributer is empty in row ' + row_number + '\n')
    else:
        contributer = str(r[8].value)
    return contributer


# gets the date
def get_date(r, row_number):
    if r[2].value is None:
        logf.write('There is no date in row ' + row_number + '\n')
    else:
        get_date = str(r[2].value)
    return get_date


# does the job of clearnace checking for the first url
def scrape_result_value(url, last_name_search, first_name_search, iter_rand):
    driver.get(url)
    driver.maximize_window()  # maxout the window size
    # finds the search bars for last name and first name
    last_name = driver.find_element_by_id('ctl00_cpExclusions_txtSPLastName')
    first_name = driver.find_element_by_id('ctl00_cpExclusions_txtSPFirstName')
    last_name.send_keys(last_name_search)
    first_name.send_keys(first_name_search)
    last_name.send_keys(Keys.ENTER)
    # first_name.send_keys(Keys.ENTER)

    # # scrolls the window
    # driver.execute_script("window.scrollTo(0, 100)")
    # # element = driver.find_element_by_id('content')
    # #
    # driver.save_screenshot(
    #     ".\\Outputs_" + iter_rand + "\\Screenshots_" + iter_rand + "\\" + last_name_search + "_" + first_name_search + ".png")
    # opening up connection, grabbing the search results page
    page_html = driver.page_source

    # driver.close()
    # lxml parsing through the current results page
    page_soup = soup(page_html, "lxml")

    # Name = page_soup.find("div", {"id": "ctl00_cpExclusions_pnlEmpty"}).ul.li.text

    # gets the no results text
    # if "No Results" is in the page then return no results variable
    noe = page_soup.find("div", {"id": "ctl00_cpExclusions_pnlEmpty"})
    search_conducted = page_soup.find("div", {"class": "timeStampResults"})
    # page_soup.find('span', {'id':"ctl00_cpExclusions_lblDatabaseDateTime" })
    search_timestamp = search_conducted.p.text

    print(f"Task: Search conducted for {first_name_search} {last_name_search}")
    print(f"Task: Scraping results for {first_name_search} {last_name_search}")
    if noe is not None:
        # scrolls the window
        driver.execute_script("window.scrollTo(0, 100)")
        # element = driver.find_element_by_id('content')
        #
        driver.save_screenshot(
            ".\\Outputs_" + iter_rand + "\\Screenshots_" + iter_rand + "\\" + last_name_search + "_" + first_name_search + ".png")
        no_results = "No Results"
    else:  # else then that means there is results for the person so scrape the info of the person
        # gets the rows
        driver.execute_script("window.scrollTo(0, 300)")
        driver.save_screenshot(
            ".\\Outputs_" + iter_rand + "\\Screenshots_" + iter_rand + "\\" + last_name_search + "_" + first_name_search + ".png")

        rows = page_soup.find("table", {"class": "leie_search_results"}).find("tbody").findAll("tr")

        # iterate through each row
        for row in rows:
            cols = row.findAll("td")
            for col in cols:
                cell = col.text
        if cell.find(last_name_search.upper()):
            no_results = "Yes"

    return no_results, search_timestamp, page_html


# checks if the person exists in the gov data base.
def clr_check(c_check, r, url_col):
    # for word in c_check.split():
    c2 = r[url_col]  # should switch to the row number
    # print(c_check)
    if c_check == 'No Results':
        no_results_cell = 'No, individual is not listed'
        c2.value = no_results_cell
    elif c_check == 'Yes':
        c2.value = 'Yes, individual appears on this list'

    return c2.value


def insert_date(r, date_col):
    d = datetime.datetime.today()
    r[date_col].value = d
    r[date_col + 1].value = d + timedelta(days=366)
    return r[date_col].value, r[date_col + 1].value


# fills not listed autrhos in nex availble row in b sheet
def set_not_listed_sheet(r, b_sheet):
    if str(r[6].value) == 'Yes, individual appears on this list':
        b_sheet.cell(row=b_sheet.max_row + 1, column=1).value = str(r[0].value)
        for num in range(2, 7):
            b_sheet.cell(row=b_sheet.max_row, column=num).value = str(r[num - 1].value)


# create_doc does the functionality does the job of automating generation of debarment word files for each author.
def create_doc(first_name_docx, last_name_docx, institution, city, state, contributer, date_checked, a_res_value,
               iter_rand, timestamp_res
               ):
    d = datetime.datetime.today()  # current date
    date_var = d.strftime("%d-%B-%Y %H:%M:%S")  # format for creation date
    print("new doc created")
    document = Document()  # generate new .doc document
    heading = document.add_heading('Debarment Check', 0) # header title
    # heading_style = document.styles['Heading 1']
    # heading_style.size = Pt(20)
    p = document.add_paragraph(
        'Prior to being invited to participate in development/authoring of a publication sponsored '
        'by Genzyme/Sanofi, a debarment check must be completed for each US author.')

    # table 1 for author name, name of institution
    records = (
        ('Author Name', first_name_docx + ' ' + last_name_docx),
        ('Name of Institution/Organization', institution),
        ('City, State', city + ', ' + state)
    )

    # data structure for table 2.
    debarment_list = (
        ('Office of Inspectors General List of Excluded Individuals.', a_res_value),
    )

    # adds table 1
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Author Information'
    hdr_cells[1].text = ' '

    # iterates through the table and the list in records data structure
    for qty, id in records:
        row_cells = table.add_row().cells
        row_cells[0].text = qty
        row_cells[1].text = id

    document.add_paragraph('')

    b_table = document.add_table(rows=1, cols=2)
    b_table.style = 'Table Grid'
    b_hdr_cells = b_table.rows[0].cells
    b_hdr_cells[0].text = 'Source for Search'
    b_hdr_cells[1].text = 'Result'

    for d_list, findings in debarment_list:
        b_row_cells = b_table.add_row().cells
        b_row_cells[0].text = d_list
        b_row_cells[1].text = findings

    document.add_paragraph(
        '\nIf the potential author is listed on any of the above, they may not be invited to author a '
        'publication sponsored by Genzyme or Sanofi; advise publication lead of findings of this '
        'search.')
    document.add_paragraph(
        'Once the debarment check has been completed, upload this document to the appropriate record '
        'in Datavision.')
    completion = document.add_paragraph('')

    completion.add_run('Debarment check completed by: ' + contributer + '\n').bold = True
    # document.add_paragraph()
    completion.add_run('Date check completed: ' + date_var + '\n').bold = True

    document.add_picture(
        '.\\Outputs_' + iter_rand + '\\Screenshots_' + iter_rand + '\\' + last_name_docx + '_' + first_name_docx + '.png',
        width=Inches(6))
    document.add_paragraph(timestamp_res)
    # Set a cell background (shading) color to RGB D9D9D9.
    a_cell_1 = table.cell(0, 0)
    a_co = a_cell_1._tc.get_or_add_tcPr()
    a_cell_2 = table.cell(0, 1)
    a_ct = a_cell_2._tc.get_or_add_tcPr()

    a_cell_color_1 = OxmlElement('w:shd')
    a_cell_color_1.set(qn('w:fill'), '#E6E6FA')

    a_cell_color_2 = OxmlElement('w:shd')
    a_cell_color_2.set(qn('w:fill'), '#E6E6FA')

    a_co.append(a_cell_color_1)
    a_ct.append(a_cell_color_2)

    b_cell_1 = b_table.cell(0, 0)
    b_co = b_cell_1._tc.get_or_add_tcPr()
    b_cell_2 = b_table.cell(0, 1)
    b_ct = b_cell_2._tc.get_or_add_tcPr()

    b_cell_color_1 = OxmlElement('w:shd')
    b_cell_color_1.set(qn('w:fill'), '#E6E6FA')

    b_cell_color_2 = OxmlElement('w:shd')
    b_cell_color_2.set(qn('w:fill'), '#E6E6FA')

    b_co.append(b_cell_color_1)
    b_ct.append(b_cell_color_2)

    if a_res_value == 'No, individual is not listed':
        document.save(
            '.\\Outputs_' + iter_rand + '\\Debarment_files_' + iter_rand + '\\' + last_name_docx + '_' + first_name_docx + '_' + d.strftime(
                '%d_%m_%Y') + '.docx')

    else:
        document.save(
            '.\\Outputs_' + iter_rand + '\\Flagged_Authors_' + iter_rand + '\\' + last_name_docx + '_' + first_name_docx + '_' + d.strftime(
                '%d_%m_%Y') + '.docx')


# all excel processing
def gener_tasks(iter_rand):
    for i, j in enumerate(sheet.iter_rows()):
        if i == 0:
            continue


        last_row = sheet.max_row
        while sheet.cell(column=1, row=last_row).value is None and last_row > 0:
            last_row -= 1
        last_col_a_value = sheet.cell(column=1, row=last_row).value
        create_sheet()
        # gets the last name first name string
        row_number = str(i + 1)
        print('Row Number: ' + row_number)
        l, f = get_name(j)

        if l == "None" or l is None:
            continue

        # gets the scraped data of result value and timestamp and does the screenshot
        chk, timestamp_results, new_url = scrape_result_value(r_url, l, f, iter_rand)
        clr_check(chk, j, 7)  # store the result to the current row and result column excel
        insert_date(j, 2)
        b_sheet = create_sheet()
        set_not_listed_sheet(j, b_sheet)  # sets the not listed dhseet
        c_cell = str(sheet['G' + str(i)].value)


        create_doc(f, l, get_institution(j, row_number), get_city(j, row_number), get_state(j, row_number), get_contributer(j, row_number),
                   get_date(j, row_number), str(j[7].value),
                   iter_rand, timestamp_results)
        print(f'Task: files created for {f} {l}')
        done = True
        ex.save(ex_name)


# fills not listed autrhos in nex availble row in b sheet
def missed_list_sheet(r, c_sheet):
    c_sheet.cell(row=c_sheet.max_row + 1, column=1).value = str(r[0].value)
    for num in range(2, 7):
        c_sheet.cell(row=c_sheet.max_row, column=num).value = str(r[num - 1].value)


def if_file_created(names_array):

    for i, j in enumerate(sheet.iter_rows()):
        if i == 0:
            continue

        print(i)

        # gets the last name first name string
        l, f = get_name(j)

        if l not in names_array:
            print('file not created')
            c_sheet = create_sheet_c()
            missed_list_sheet(j, c_sheet)

            # insert that row
        else:
            print('file Created')

            # b_sheet = create_sheet()
            # set_not_listed_sheet(j, b_sheet)  # sets the not listed dhseet

            # ex.save('db_check.xlsx')


def string_man(file):
    return os.path.splitext(file)[0]


def execute_process():
    filename_with_ext = str(get_working_filename())
    with_ext = os.path.basename(filename_with_ext)
    print(with_ext)  # name of the file
    filename = string_man(with_ext)
    print(f"{filename}")
    rand = str(random.random())
    # if os.path.exists(output) is False:
    directory_structure(filename)  # creating all directory structure

    gener_tasks(filename)  # starts the for each row

    last_array = check_not_missed(filename)  # array of file name

    print(if_file_created(last_array))
    ex.save(ex_name)

    if done:
        mv_dir_structure(filename)  # moving the database to the completed directory
    else:
        print("the file processing is not done")


start_time = time.time()
execute_process()
print("--- %s seconds ---" % (time.time() - start_time))
print("closing connection")
# driver.close()
print("Process completed. You may exit the window.")
