#!/bin/python
from bs4 import BeautifulSoup as soup
import os
from docx import Document
from docx.shared import Inches
import openpyxl
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from urllib.request import urlopen as uReq
from bs4 import BeautifulSoup as soup
from docx.oxml.shared import OxmlElement, qn
import datetime
import random
from datetime import timedelta
import shutil

Root_path = 'C:\\Users\\yangb\\PycharmProjects\\DebarmentCheckAIBot'
# change directory
os.chdir('C:\\Users\\yangb\\PycharmProjects\\DebarmentCheckAIBot')
logf = open(Root_path + 'log_file.txt', 'w')


# define the name of the directory to be created
# directory structure for screenshots, files, and compeleted files
def directory_structure(iter_rand):
    output = Root_path + 'Outputs_' + iter_rand
    flagged_authors = output + '\\Flagged_authors_files_' + iter_rand
    debarment_file_path = output + '\\Debarment_files_' + iter_rand
    screenshots_path = output + '\\Screenshots_' + iter_rand
    completed_file_path = output + '\\Completed_file_' + iter_rand

    # error handlers for creating directory for screenshots and debarment file
    try:
        os.mkdir(output)
    except OSError as e:
        print("Creation of the directory %s failed" % output)
        logf.write('creation of the directory {0}:{1} failed\n'.format(str(output), str(e)))
        # print("Creation of the directory %s failed" % screenshots_path)
    else:
        print("Successfully created the directory %s " % output)

    try:
        os.mkdir(flagged_authors)
    except OSError as e:
        print("Creation of the directory %s failed" % flagged_authors)
        logf.write('creation of the directory {0}:{1} failed\n'.format(str(flagged_authors), str(e)))
        # print("Creation of the directory %s failed" % screenshots_path)
    else:
        print("Successfully created the directory %s " % flagged_authors)

    try:
        os.mkdir(screenshots_path)
    except OSError as e:
        print(e)
        logf.write('creation of the directory {0}:{1} failed\n'.format(str(screenshots_path), str(e)))
        # print("Creation of the directory %s failed" % screenshots_path)
    else:
        print("Successfully created the directory %s " % screenshots_path)

    try:
        os.mkdir(debarment_file_path)
    except OSError:
        print("Creation of the directory %s failed" % debarment_file_path)
        logf.write('creation of the directory {0}:{1} failed\n'.format(str(debarment_file_path), str(e)))
    else:
        print("Successfully created the directory %s " % debarment_file_path)

    try:
        os.mkdir(completed_file_path)
    except OSError:
        print("Creation of the directory %s failed" % completed_file_path)
        logf.write('creation of the directory {0}:{1} failed\n'.format(str(completed_file_path), str(e)))
    else:
        print("Successfully created the directory %s " % completed_file_path)


# gets the input variable for the process execution (excel file, list of authors)
def get_working_filename():
    if len(os.listdir('.\\working_dir')) == 0:
        logf.write('empty directory\n')
    elif len(os.listdir('.\\working_dir')) > 1:
        logf.write('too many files, one directory at a time\n')
    else:
        for file in os.listdir('.\\working_dir'):  # iterate through all the files
            if file.endswith('.xlsx'):  # identify if the file ends with.xlsx
                move_file = os.path.join(".\\working_dir", file)  # gets the file in working directory
                return move_file
            else:  # if file not ends with .xlsx then its not an excel file
                logf.write('not an excel file\n')


def check_not_missed(iter_rand):
    array = []
    for file in os.listdir(
            '.\\Outputs_' + iter_rand + '\\Debarment_files' + iter_rand):  # iterate through all the files
        split_filename = file.split('_')
        array.append(split_filename[0])  # gets only the last name
    return array


# get file and move file
def mv_dir_structure(iter_rand):
    try:
        # move the current file to the completed directory
        shutil.move(ex_name, ".\\Outputs_" + iter_rand + "\\Completed_files" + iter_rand)
    except shutil.Error as err:
        logf.write(str(err) + '\n')
        # errors.extend(err.args[0])


# get webdriver for chrome chromedriver.exe path - this would have to change for everyone
driver = webdriver.Chrome('C:\\Users\\yangb\\Desktop\\chromedriver.exe')
driver.maximize_window()  # maxout the window size

# sets the window size. This size is specified from screenshot reasons
# driver.set_window_size(1100, 1500)  # ideal was 1100, 1500
done = True  # flagger variable for flagging the process
# getting the file from the working_dir
ex_name = get_working_filename()  # excel file name (database)
ex = openpyxl.load_workbook(ex_name)  # opens the excel
sheet = ex['Sheet1']  # gets the first tab sheet

# url 1 for search
r_url = 'https://exclusions.oig.hhs.gov/'


# new sheet for not listed people
# b_sheet.cell(row=b_sheet.max_row, column=1).value = 'Not on the list'
def create_sheet():
    if 'Flagged' not in ex.sheetnames:
        ex.create_sheet('Flagged')
        ex.save(ex_name)
    b_sheet = ex['Flagged']
    return b_sheet


def create_sheet_c():
    if 'Missed' not in ex.sheetnames:
        ex.create_sheet('Missed')
        ex.save(ex_name)
    c_sheet = ex['Missed']
    return c_sheet


# get name of one row
# same function
def get_name(r):
    ex_last = str(r[0].value).replace(" ", "")
    ex_first = str(r[1].value).replace(" ", "")

    return ex_last, ex_first


# gets institution
def get_institution(r, row_number):
    if r[4].value is None:
        logf.write('institution is empty in row ' + row_number + '\n')
    else:
        institute = str(r[4].value)
    return institute


# gets the city state
def get_city_state(r, row_number):
    if r[5].value is None:
        logf.write('institution is empty in row ' + row_number + '\n')
    else:
        location = str(r[5].value)
    return location


# gets the name of the contributer
def get_contributer(r, row_number):
    if r[7].value is None:
        logf.write('contributer is empty in row ' + row_number + '\n')
    else:
        contributer = str(r[7].value)
    return contributer


# gets the date
def get_date(r, row_number):
    if r[2].value is None:
        logf.write('There is no date in row ' + row_number + '\n')

    else:
        get_date = str(r[2].value)
    return get_date


# does the job of clearnace checking for the first url
def scrape_result_value(url, last_name_search, first_name_search, iter_rand):
    driver.get(url)
    # finds the search bars for last name and first name
    last_name = driver.find_element_by_id('ctl00_cpExclusions_txtSPLastName')
    first_name = driver.find_element_by_id('ctl00_cpExclusions_txtSPFirstName')
    last_name.send_keys(last_name_search)
    first_name.send_keys(first_name_search)
    last_name.send_keys(Keys.ENTER)
    # first_name.send_keys(Keys.ENTER)

    # scrolls the window
    driver.execute_script("window.scrollTo(0, 100)")
    # element = driver.find_element_by_id('content')
    #
    driver.save_screenshot(
        ".\\Outputs_" + iter_rand + "\\Screenshots" + iter_rand + "\\" + last_name_search + "_" + first_name_search + ".png")
    # opening up connection, grabbing the search results page
    page_html = driver.page_source

    # driver.close()
    # lxml parsing through the current results page
    page_soup = soup(page_html, "lxml")

    # Name = page_soup.find("div", {"id": "ctl00_cpExclusions_pnlEmpty"}).ul.li.text

    # gets the no results text
    # if "No Results" is in the page then return no results variable
    noe = page_soup.find("div", {"id": "ctl00_cpExclusions_pnlEmpty"})
    search_conducted = page_soup.find("div", {"class": "timeStampResults"})
    search_timestamp = search_conducted.p.text
    print(search_timestamp)
    if noe is not None:
        no_results = "No Results"
    else:  # else then that means there is results for the person so scrape the info of the person
        # gets the rows
        rows = page_soup.find("table", {"class": "leie_search_results"}).find("tbody").findAll("tr")

        # iterate through each row
        for row in rows:
            cols = row.findAll("td")
            for col in cols:
                cell = col.text
        if cell.find(last_name_search.upper()):
            no_results = "Yes"

    return no_results, search_timestamp, page_html


# checks if the person exists in the gov data base.
def clr_check(c_check, r, url_col):
    # for word in c_check.split():
    c2 = r[url_col]  # should switch to the row number
    # print(c_check)
    if c_check == 'No Results':
        no_results_cell = 'No, individual is not listed'
        c2.value = no_results_cell
    elif c_check == 'Yes':
        c2.value = 'Yes, individual appears on this list'

    return c2.value


def insert_date(r, date_col):
    d = datetime.datetime.today()
    r[date_col].value = d
    r[date_col + 1].value = d + timedelta(days=366)
    return r[date_col].value, r[date_col + 1].value


# fills not listed autrhos in nex availble row in b sheet
def set_not_listed_sheet(r, b_sheet):
    if str(r[6].value) == 'Yes, individual appears on this list':
        b_sheet.cell(row=b_sheet.max_row + 1, column=1).value = str(r[0].value)
        for num in range(2, 7):
            b_sheet.cell(row=b_sheet.max_row, column=num).value = str(r[num - 1].value)


# create_doc does the functionality does the job of automating generation of debarment word files for each author.
def create_doc(first_name_docx, last_name_docx, institution, city_state, contributer, date_checked, a_res_value,
               iter_rand, timestamp_res
               ):
    d = datetime.datetime.today()  # current date
    date_var = d.strftime("%d-%B-%Y %H:%M:%S")  # format for creation date

    document = Document()  # generate new .doc document
    document.add_heading('Debarment Check', 0)  # header title
    p = document.add_paragraph(
        'Prior to being invited to participate in development/authoring of a publication sponsored '
        'by Genzyme/Sanofi, a debarment check must be completed for each US author.')

    # table 1 for author name, name of institution
    records = (
        ('Author Name', first_name_docx + ' ' + last_name_docx),
        ('Name of Institution', institution),
        ('City, State', city_state)
    )

    # data structure for table 2.
    debarment_list = (
        ('Office of Inspectors General List of Excluded Individuals.', a_res_value),
    )

    # adds table 1
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Information'
    hdr_cells[1].text = 'Id'

    # iterates through the table and the list in records data structure
    for qty, id in records:
        row_cells = table.add_row().cells
        row_cells[0].text = qty
        row_cells[1].text = id

    document.add_paragraph('')

    b_table = document.add_table(rows=1, cols=2)
    b_table.style = 'Table Grid'
    b_hdr_cells = b_table.rows[0].cells
    b_hdr_cells[0].text = 'Debarment List'
    b_hdr_cells[1].text = 'Findings'

    for d_list, findings in debarment_list:
        b_row_cells = b_table.add_row().cells
        b_row_cells[0].text = d_list
        b_row_cells[1].text = findings

    document.add_paragraph(
        '\nIf the potential author is listed on any of the above, they may not be invited to author a '
        'publication sponsored by Genzyme or Sanofi; advise publication lead of findings of this '
        'search.')
    document.add_paragraph(
        'Once the debarment check has been completed, upload this document to the appropriate record '
        'in Datavision.')
    completion = document.add_paragraph('')

    completion.add_run('Debarment check completed by: ' + contributer + '\n').bold = True
    # document.add_paragraph()
    completion.add_run('Date check completed: ' + date_var + '\n').bold = True

    document.add_picture(
        '.\\Outputs_' + iter_rand + '\\Screenshots' + iter_rand + '\\' + last_name_docx + '_' + first_name_docx + '.png',
        width=Inches(6))
    document.add_paragraph(timestamp_res)
    # Set a cell background (shading) color to RGB D9D9D9.
    a_cell_1 = table.cell(0, 0)
    a_co = a_cell_1._tc.get_or_add_tcPr()
    a_cell_2 = table.cell(0, 1)
    a_ct = a_cell_2._tc.get_or_add_tcPr()

    a_cell_color_1 = OxmlElement('w:shd')
    a_cell_color_1.set(qn('w:fill'), '#94C167')

    a_cell_color_2 = OxmlElement('w:shd')
    a_cell_color_2.set(qn('w:fill'), '#94C167')

    a_co.append(a_cell_color_1)
    a_ct.append(a_cell_color_2)

    b_cell_1 = b_table.cell(0, 0)
    b_co = b_cell_1._tc.get_or_add_tcPr()
    b_cell_2 = b_table.cell(0, 1)
    b_ct = b_cell_2._tc.get_or_add_tcPr()

    b_cell_color_1 = OxmlElement('w:shd')
    b_cell_color_1.set(qn('w:fill'), '#94C167')

    b_cell_color_2 = OxmlElement('w:shd')
    b_cell_color_2.set(qn('w:fill'), '#94C167')

    b_co.append(b_cell_color_1)
    b_ct.append(b_cell_color_2)

    if a_res_value == 'No, individual is not listed':
        document.save(
            '.\\Outputs_' + iter_rand + '\\Debarment_files' + iter_rand + '\\' + last_name_docx + '_' + first_name_docx + '_' + d.strftime(
                '%d_%m_%Y') + '.docx')

    else:
        document.save(
            '.\\Outputs_' + iter_rand + '\\Flagged_Authors' + iter_rand + '\\' + last_name_docx + '_' + first_name_docx + '_' + d.strftime(
                '%d_%m_%Y') + '.docx')


# all excel processing
def gener_tasks(iter_rand):
    for i, j in enumerate(sheet.iter_rows()):
        if i == 0:
            continue

        print(i)
        create_sheet()
        # gets the last name first name string
        l, f = get_name(j)
        # gets the scraped data of result value and timestamp and does the screenshot
        chk, timestamp_results, new_url = scrape_result_value(r_url, l, f, iter_rand)
        clr_check(chk, j, 6)  # store the result to the current row and result column excel
        insert_date(j, 2)
        b_sheet = create_sheet()
        set_not_listed_sheet(j, b_sheet)  # sets the not listed dhseet
        c_cell = str(sheet['G' + str(i)].value)
        row_number = str(i + 1)

        create_doc(f, l, get_institution(j, row_number), get_city_state(j, row_number), get_contributer(j, row_number),
                   get_date(j, row_number), str(j[6].value),
                   iter_rand, timestamp_results)
        done = True
        ex.save('db_check.xlsx')


# fills not listed autrhos in nex availble row in b sheet
def missed_list_sheet(r, c_sheet):

    c_sheet.cell(row=c_sheet.max_row + 1, column=1).value = str(r[0].value)
    for num in range(2, 7):
        c_sheet.cell(row=c_sheet.max_row, column=num).value = str(r[num - 1].value)
ex.save(ex_name)


def if_file_created(names_array):
    for i, j in enumerate(sheet.iter_rows()):
        if i == 0:
            continue

        print(i)

        # gets the last name first name string
        l, f = get_name(j)

        if l not in names_array:
            print('file not created')
            c_sheet = create_sheet_c()
            missed_list_sheet(j, c_sheet)

            # insert that row
        else:
            print('file Created')

            # b_sheet = create_sheet()
            # set_not_listed_sheet(j, b_sheet)  # sets the not listed dhseet

            # ex.save('db_check.xlsx')


def string_man(file):
    return os.path.splitext(file)[0]


def execute_process():
    filename_with_ext = str(get_working_filename())
    with_ext = os.path.basename(filename_with_ext)
    print(with_ext)  # name of the file
    filename = string_man(with_ext)
    print(filename)
    rand = str(random.random())
    directory_structure(filename)  # creating all directory structure
    gener_tasks(filename)  # starts the for each row

    last_array = check_not_missed(filename)  # array of file name

    print(if_file_created(last_array))
    ex.save(ex_name)

    if done:
        mv_dir_structure(filename)  # moving the database to the completed directory
    else:
        print("the file processing is not done")


execute_process()
driver.close()
